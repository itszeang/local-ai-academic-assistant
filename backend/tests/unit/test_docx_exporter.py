from pathlib import Path

from docx import Document as DocxDocument

from app.common.types import GenerationMode, MISSING_INFORMATION_RESPONSE
from app.export.docx_exporter import DocxExporter


def test_docx_exporter_preserves_sections_and_references(tmp_path: Path) -> None:
    output_path = tmp_path / "supported.docx"

    DocxExporter().export(
        output={
            "id": "output_1",
            "mode": GenerationMode.QA.value,
            "title": "Grounded Answer",
            "sections": [
                {"heading": "Answer", "blocks": ["Motivation improves persistence. (Aksoy, 2024)"]},
                {"heading": "Supporting Evidence", "blocks": ["The source states this directly."]},
            ],
            "references": ["motivation.pdf, p. 5"],
            "fallback_used": False,
        },
        path=output_path,
    )

    doc = DocxDocument(output_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    assert "Grounded Answer" in paragraphs
    assert "Answer" in paragraphs
    assert "Motivation improves persistence. (Aksoy, 2024)" in paragraphs
    assert "References" in paragraphs
    assert "motivation.pdf, p. 5" in paragraphs


def test_docx_exporter_preserves_fallback_without_references(tmp_path: Path) -> None:
    output_path = tmp_path / "fallback.docx"

    DocxExporter().export(
        output={
            "id": "output_1",
            "mode": GenerationMode.QA.value,
            "title": MISSING_INFORMATION_RESPONSE,
            "sections": [{"heading": "Answer", "blocks": [MISSING_INFORMATION_RESPONSE]}],
            "references": [],
            "fallback_used": True,
        },
        path=output_path,
    )

    doc = DocxDocument(output_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    assert paragraphs.count(MISSING_INFORMATION_RESPONSE) >= 2
    assert "References" not in paragraphs
