"""DOCX export for structured academic outputs."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from app.common.types import AcademicOutput


class DocxExporter:
    def export(self, *, output: AcademicOutput | dict[str, Any], path: Path) -> Path:
        payload = self._payload(output)
        export_path = Path(path)
        export_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(str(payload["title"]), level=1)

        for section in payload["sections"]:
            document.add_heading(str(section["heading"]), level=2)
            for block in section["blocks"]:
                document.add_paragraph(str(block))

        references = payload["references"]
        if references:
            document.add_heading("References", level=2)
            for reference in references:
                document.add_paragraph(str(reference))

        document.save(export_path)
        return export_path

    @staticmethod
    def _payload(output: AcademicOutput | dict[str, Any]) -> dict[str, Any]:
        if isinstance(output, AcademicOutput):
            return {
                "id": output.id,
                "mode": output.mode.value,
                "title": output.title,
                "sections": output.sections,
                "references": output.references,
                "fallback_used": output.fallback_used,
            }
        return output
